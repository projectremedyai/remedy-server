"""office-verify deterministic rule engine.

Each rule is a pure function ``(DocxContext) -> OfficeCheckResult`` registered
in ``DOCX_RULES`` under its canonical rule id. Determinism contract (NFR1):
no network, no LLM, no clock — same input bytes, same report.
"""

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Callable

from project_remedy.models import FileType
from project_remedy.office_acceptance import (
    OfficeCheckReport,
    OfficeCheckResult,
    _docx_outline_level,
    _docx_paragraph_has_heading_structure,
    _infer_file_type,
)
from project_remedy.office_rules import NS, RULE_CATALOG, RULE_SPECS_BY_ID, qn_w


@dataclass
class DocxContext:
    """Everything a docx rule may read, loaded exactly once per document."""

    path: Path
    document: Any                 # python-docx Document
    body_root: ET.Element         # parsed word/document.xml

    @classmethod
    def load(cls, path: Path) -> "DocxContext":
        from docx import Document

        with zipfile.ZipFile(path) as zf:
            body_root = ET.fromstring(zf.read("word/document.xml"))
        return cls(path=Path(path), document=Document(str(path)), body_root=body_root)


DOCX_RULES: dict[str, Callable[[DocxContext], OfficeCheckResult]] = {}


def docx_rule(rule_id: str):
    def wrap(fn: Callable[[DocxContext], OfficeCheckResult]):
        DOCX_RULES[rule_id] = fn
        return fn
    return wrap


def _make_result(rule_id: str, *, flagged: bool, details: list[str]) -> OfficeCheckResult:
    spec = RULE_SPECS_BY_ID[rule_id]
    return OfficeCheckResult(
        rule_id=spec.emitted_id,
        description=spec.description,
        status=spec.flag_status if flagged else "Passed",
        details=details if flagged else [],
        fixable=spec.fixable,
        checkpoint=spec.checkpoint,
        wcag_ref=spec.wcag_ref,
    )


# --- Checkpoint 1: document metadata ---------------------------------------

@docx_rule("OOXML-DOCX-1.1")
def rule_docx_title(ctx: DocxContext) -> OfficeCheckResult:
    title = (ctx.document.core_properties.title or "").strip()
    return _make_result("OOXML-DOCX-1.1", flagged=not title,
                        details=["docProps/core.xml dc:title is empty"])


@docx_rule("OOXML-DOCX-1.2")
def rule_docx_language(ctx: DocxContext) -> OfficeCheckResult:
    language = (getattr(ctx.document.core_properties, "language", "") or "").strip()
    return _make_result("OOXML-DOCX-1.2", flagged=not language,
                        details=["docProps/core.xml dc:language is empty"])


# --- Checkpoint 2: heading structure ----------------------------------------

_HEADING_STYLE_RE = re.compile(r"^(?:accessibility )?heading\s*(\d+)?", re.IGNORECASE)


def _heading_level(paragraph: Any) -> int | None:
    """1-based heading level, or None if the paragraph is not a heading.

    Title (and Accessibility Title) count as level 1; "Heading N" is level N;
    a bare w:outlineLvl value v maps to level v+1.
    """
    style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip()
    lowered = style_name.lower()
    if lowered.startswith(("title", "accessibility title")):
        return 1
    match = _HEADING_STYLE_RE.match(lowered)
    if match:
        return int(match.group(1) or 1)
    outline = _docx_outline_level(paragraph)
    if outline is not None:
        return outline + 1
    return None


@docx_rule("OOXML-DOCX-2.1")
def rule_docx_headings_present(ctx: DocxContext) -> OfficeCheckResult:
    has_heading = any(_docx_paragraph_has_heading_structure(p) for p in ctx.document.paragraphs)
    return _make_result("OOXML-DOCX-2.1", flagged=not has_heading,
                        details=["no paragraph carries a Heading/Title style or w:outlineLvl"])


@docx_rule("OOXML-DOCX-2.2")
def rule_docx_heading_skips(ctx: DocxContext) -> OfficeCheckResult:
    skips: list[str] = []
    previous = 0  # virtual document root: the first heading must be level 1
    for index, paragraph in enumerate(ctx.document.paragraphs):
        level = _heading_level(paragraph)
        if level is None:
            continue
        if level > previous + 1:
            snippet = paragraph.text.strip()[:48]
            skips.append(f"paragraph {index}: heading level jumps {previous} -> {level} ('{snippet}')")
        previous = level
    return _make_result("OOXML-DOCX-2.2", flagged=bool(skips), details=skips)


@docx_rule("OOXML-DOCX-2.3")
def rule_docx_no_orphan_intro(ctx: DocxContext) -> OfficeCheckResult:
    first = next((p for p in ctx.document.paragraphs if p.text.strip()), None)
    if first is None:
        return _make_result("OOXML-DOCX-2.3", flagged=False, details=[])
    flagged = not _docx_paragraph_has_heading_structure(first)
    return _make_result(
        "OOXML-DOCX-2.3", flagged=flagged,
        details=[f"document opens on body paragraph: '{first.text.strip()[:64]}'"],
    )


# --- Checkpoint 3: images ----------------------------------------------------

_PLACEHOLDER_ALT_RE = re.compile(
    r"^\s*(?:image|picture|img|graphic|photo)[\s_-]*\d*\s*$"
    r"|^[\w\s_-]*\.(?:png|jpe?g|gif|bmp|tiff?|webp)\s*$",
    re.IGNORECASE
)


def _iter_image_doc_prs(body_root: ET.Element) -> list[tuple[str, ET.Element]]:
    """(kind, wp:docPr) for every inline and anchored image in body order."""
    found: list[tuple[str, ET.Element]] = []
    for drawing in body_root.iter(qn_w("drawing")):
        for kind in ("inline", "anchor"):
            for container in drawing.iter(f"{{{NS['wp']}}}{kind}"):
                doc_pr = container.find(f"{{{NS['wp']}}}docPr")
                if doc_pr is not None:
                    found.append(("inline" if kind == "inline" else "anchored", doc_pr))
    return found


def _alt_text_of(doc_pr: ET.Element) -> str:
    return ((doc_pr.get("descr") or "").strip() or (doc_pr.get("title") or "").strip())


@docx_rule("OOXML-DOCX-3.1")
def rule_docx_image_alt_present(ctx: DocxContext) -> OfficeCheckResult:
    missing: list[str] = []
    for ordinal, (kind, doc_pr) in enumerate(_iter_image_doc_prs(ctx.body_root), start=1):
        if not _alt_text_of(doc_pr):
            missing.append(f"image {ordinal} ({kind}) has no descr/title alt text")
    return _make_result("OOXML-DOCX-3.1", flagged=bool(missing), details=missing)


@docx_rule("OOXML-DOCX-3.2")
def rule_docx_alt_not_placeholder(ctx: DocxContext) -> OfficeCheckResult:
    offenders: list[str] = []
    for ordinal, (kind, doc_pr) in enumerate(_iter_image_doc_prs(ctx.body_root), start=1):
        alt = _alt_text_of(doc_pr)
        if alt and _PLACEHOLDER_ALT_RE.match(alt):
            offenders.append(f"image {ordinal} ({kind}) alt is a placeholder: '{alt}'")
    return _make_result("OOXML-DOCX-3.2", flagged=bool(offenders), details=offenders)


# --- Checkpoint 4: tables ----------------------------------------------------

@docx_rule("OOXML-DOCX-4.1")
def rule_docx_table_header_marked(ctx: DocxContext) -> OfficeCheckResult:
    unmarked: list[str] = []
    for index, tbl in enumerate(ctx.body_root.iter(qn_w("tbl")), start=1):
        first_tr = tbl.find(qn_w("tr"))
        if first_tr is None:
            continue
        tr_pr = first_tr.find(qn_w("trPr"))
        tbl_header = tr_pr.find(qn_w("tblHeader")) if tr_pr is not None else None
        # <w:tblHeader w:val="0"/> explicitly negates the header mark (ECMA-376
        # ST_OnOff); a bare element or val="1"/"true"/"on" means marked.
        header_val = (tbl_header.get(qn_w("val")) or "").lower() if tbl_header is not None else None
        if tbl_header is None or header_val in ("0", "false", "off"):
            unmarked.append(f"table {index}: first row lacks an effective w:tblHeader")
    return _make_result("OOXML-DOCX-4.1", flagged=bool(unmarked), details=unmarked)


@docx_rule("OOXML-DOCX-4.2")
def rule_docx_no_merged_header_cells(ctx: DocxContext) -> OfficeCheckResult:
    merged: list[str] = []
    for index, tbl in enumerate(ctx.body_root.iter(qn_w("tbl")), start=1):
        first_tr = tbl.find(qn_w("tr"))
        if first_tr is None:
            continue
        for tc_pr in first_tr.iter(qn_w("tcPr")):
            if tc_pr.find(qn_w("gridSpan")) is not None or tc_pr.find(qn_w("vMerge")) is not None:
                merged.append(f"table {index}: header row contains merged cells (gridSpan/vMerge)")
                break
    return _make_result("OOXML-DOCX-4.2", flagged=bool(merged), details=merged)


# --- Checkpoints 5-7: lists, hyperlinks, color-only meaning ------------------

_MANUAL_BULLET_RE = re.compile(r"^\s*(?:[•\-\*]\s+|\d+[.)]\s+)")
_BARE_URL_RE = re.compile(r"(?i)^\s*(?:https?://|www\.)\S+\s*$")
_GENERIC_LINK_TEXT = {"click here", "here", "read more", "more", "learn more", "link", "this link"}
_COLOR_PHRASE_RE = re.compile(
    r"\b(?:in|shown in|marked in|highlighted in|displayed in)\s+"
    r"(?:red|green|blue|yellow|orange|purple|pink)\b"
    r"|\b(?:red|green|blue|yellow|orange|purple|pink)\s+(?:text|items?|entries|values|fields?|cells?)\b",
    re.IGNORECASE
)


def _p_text(p_element: ET.Element) -> str:
    return "".join(t.text or "" for t in p_element.iter(qn_w("t")))


@docx_rule("OOXML-DOCX-5.1")
def rule_docx_manual_bullets(ctx: DocxContext) -> OfficeCheckResult:
    flagged: list[str] = []
    for index, p in enumerate(ctx.body_root.iter(qn_w("p")), start=1):
        text = _p_text(p)
        if not _MANUAL_BULLET_RE.match(text):
            continue
        p_pr = p.find(qn_w("pPr"))
        has_num_pr = p_pr is not None and p_pr.find(qn_w("numPr")) is not None
        if not has_num_pr:
            flagged.append(f"paragraph {index}: manual bullet/number without w:numPr: '{text.strip()[:48]}'")
    return _make_result("OOXML-DOCX-5.1", flagged=bool(flagged), details=flagged)


@docx_rule("OOXML-DOCX-6.1")
def rule_docx_link_text(ctx: DocxContext) -> OfficeCheckResult:
    offenders: list[str] = []
    for index, link in enumerate(ctx.body_root.iter(qn_w("hyperlink")), start=1):
        display = "".join(t.text or "" for t in link.iter(qn_w("t"))).strip()
        if not display:
            continue
        if _BARE_URL_RE.match(display) or display.lower() in _GENERIC_LINK_TEXT:
            offenders.append(f"hyperlink {index}: display text is not descriptive: '{display[:64]}'")
    return _make_result("OOXML-DOCX-6.1", flagged=bool(offenders), details=offenders)


@docx_rule("OOXML-DOCX-7.1")
def rule_docx_color_only_meaning(ctx: DocxContext) -> OfficeCheckResult:
    flagged: list[str] = []
    for index, p in enumerate(ctx.body_root.iter(qn_w("p")), start=1):
        text = _p_text(p)
        if not text.strip() or not _COLOR_PHRASE_RE.search(text):
            continue
        has_colored_run = any(
            (color.get(qn_w("val")) or "").lower() not in ("", "auto", "000000")
            for color in p.iter(qn_w("color"))
        )
        if has_colored_run:
            flagged.append(
                f"paragraph {index}: color-reference phrase with colored run — verify meaning "
                f"is not conveyed by color alone: '{text.strip()[:64]}'"
            )
    return _make_result("OOXML-DOCX-7.1", flagged=bool(flagged), details=flagged)


# --- Engine entry point -------------------------------------------------------

class OfficeAccessibilityChecker:
    """FR1: runs the full deterministic rule catalog for one Office document.

    Phase 1 implements docx; pptx/xlsx delegate to the legacy per-format
    checks in ``office_acceptance`` until Phases 2/3 move them here.
    """

    def __init__(self, file_path: Path, file_type: FileType | None = None) -> None:
        self.file_path = Path(file_path)
        self.file_type = file_type or _infer_file_type(self.file_path)

    def run_all(self) -> OfficeCheckReport:
        if self.file_type != FileType.DOCX:
            from project_remedy.office_acceptance import _check_pptx, _check_xlsx

            if self.file_type == FileType.PPTX:
                return _check_pptx(self.file_path)
            if self.file_type == FileType.XLSX:
                return _check_xlsx(self.file_path)
            raise ValueError(f"Unsupported Office checker type: {self.file_type}")
        ctx = DocxContext.load(self.file_path)
        results = [
            DOCX_RULES[spec.rule_id](ctx)
            for spec in RULE_CATALOG
            if spec.format == "docx"
        ]
        return OfficeCheckReport(file_path=self.file_path, file_type=FileType.DOCX, results=results)
