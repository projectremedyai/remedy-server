"""Smoke tests for the in-memory OOXML fixture builders."""

from __future__ import annotations

import zipfile

from tests.unit.office_fixtures import make_docx, make_pptx, make_xlsx


def test_make_docx_produces_openable_package(tmp_path):
    path = make_docx(
        tmp_path / "sample.docx",
        title="Sample",
        language="en-US",
        headings=[("Sample", 0), ("Section One", 1)],
        body_paragraphs=["This is a body paragraph with enough words to look like prose."],
        tables=1,
        inline_images=1,
        image_alt="A sample image",
    )
    from docx import Document

    doc = Document(str(path))
    assert doc.core_properties.title == "Sample"
    assert len(doc.tables) == 1
    assert len(doc.inline_shapes) == 1


def test_make_docx_anchored_image_is_invisible_to_inline_shapes(tmp_path):
    path = make_docx(tmp_path / "anchored.docx", inline_images=1, image_alt=None, anchored_images=True)
    from docx import Document

    doc = Document(str(path))
    # the anchored conversion hides the image from python-docx's inline API —
    # exactly the baseline gap OOXML-DOCX-3.1 must close
    assert len(doc.inline_shapes) == 0
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    assert "<wp:anchor" in xml


def test_make_pptx_and_xlsx_open(tmp_path):
    pptx_path = make_pptx(tmp_path / "deck.pptx", title="Deck", slides=2, slide_titles=True, pictures=1)
    xlsx_path = make_xlsx(tmp_path / "book.xlsx", title="Book", data_rows=3, data_cols=2)
    from openpyxl import load_workbook
    from pptx import Presentation

    assert len(list(Presentation(str(pptx_path)).slides)) == 2
    wb = load_workbook(str(xlsx_path))
    assert wb.properties.title == "Book"
