"""Rule-catalog invariants + per-rule Pass/Fail tests (grows through Task 8)."""

from __future__ import annotations

from project_remedy.office_rules import NS, RULE_CATALOG, RULE_SPECS_BY_ID, RuleSpec


def test_catalog_has_twelve_docx_rules_with_unique_ids():
    docx = [s for s in RULE_CATALOG if s.format == "docx"]
    assert len(docx) == 12
    assert len({s.rule_id for s in docx}) == 12
    assert len({s.emitted_id for s in docx}) == 12


def test_every_rule_declares_xml_refs_and_wcag_ref():
    for spec in RULE_CATALOG:
        assert spec.rule_id.startswith("OOXML-DOCX-")
        assert spec.xml_refs, spec.rule_id           # FR2: self-documenting XML refs
        assert spec.wcag_ref, spec.rule_id
        assert spec.flag_status in ("Failed", "Manual Check Needed")


def test_legacy_alias_mapping_is_exact():
    aliases = {s.rule_id: s.emitted_id for s in RULE_CATALOG if s.emitted_id != s.rule_id}
    assert aliases == {
        "OOXML-DOCX-1.1": "docx-title",
        "OOXML-DOCX-1.2": "docx-language",
        "OOXML-DOCX-2.1": "docx-headings",
        "OOXML-DOCX-3.1": "docx-alt-text",
        "OOXML-DOCX-4.1": "docx-table-headers",
    }


def test_pattern_rules_route_to_manual_check_needed():
    assert RULE_SPECS_BY_ID["OOXML-DOCX-5.1"].flag_status == "Manual Check Needed"
    assert RULE_SPECS_BY_ID["OOXML-DOCX-7.1"].flag_status == "Manual Check Needed"


def test_new_rules_without_remediator_support_ship_fixable_false():
    # FR6: never fixable=True without a real office_remediator code path
    for rule_id in ("OOXML-DOCX-2.2", "OOXML-DOCX-2.3", "OOXML-DOCX-3.2",
                    "OOXML-DOCX-4.2", "OOXML-DOCX-5.1", "OOXML-DOCX-6.1", "OOXML-DOCX-7.1"):
        assert RULE_SPECS_BY_ID[rule_id].fixable is False, rule_id


def test_namespace_map_covers_wordprocessing_drawing():
    assert set(NS) >= {"w", "wp", "a", "r"}


from pathlib import Path

from project_remedy.office_checker import DOCX_RULES, DocxContext
from tests.unit.office_fixtures import make_docx


def _run(rule_id: str, path: Path):
    return DOCX_RULES[rule_id](DocxContext.load(path))


def test_rule_1_1_title(tmp_path):
    good = make_docx(tmp_path / "g.docx", title="Has Title")
    bad = make_docx(tmp_path / "b.docx", title="")
    ok = _run("OOXML-DOCX-1.1", good)
    fail = _run("OOXML-DOCX-1.1", bad)
    assert ok.status == "Passed" and ok.rule_id == "docx-title"
    assert fail.status == "Failed" and fail.fixable is True
    assert fail.checkpoint == "Document metadata" and fail.wcag_ref == "2.4.2"


def test_rule_1_2_language(tmp_path):
    good = make_docx(tmp_path / "g.docx", language="en-US")
    bad = make_docx(tmp_path / "b.docx", language="")
    assert _run("OOXML-DOCX-1.2", good).status == "Passed"
    result = _run("OOXML-DOCX-1.2", bad)
    assert result.status == "Failed" and result.rule_id == "docx-language"


def test_rule_2_1_headings_present(tmp_path):
    good = make_docx(tmp_path / "g.docx", headings=[("T", 0)])
    bad = make_docx(tmp_path / "b.docx", body_paragraphs=["Just plain body text here."])
    assert _run("OOXML-DOCX-2.1", good).status == "Passed"
    result = _run("OOXML-DOCX-2.1", bad)
    assert result.status == "Failed" and result.rule_id == "docx-headings"


def test_rule_2_2_no_level_skips(tmp_path):
    good = make_docx(tmp_path / "g.docx", headings=[("T", 0), ("A", 1), ("B", 2)])
    skip = make_docx(tmp_path / "s.docx", headings=[("A", 1), ("C", 3)])
    first_deep = make_docx(tmp_path / "f.docx", headings=[("Only", 2)])
    assert _run("OOXML-DOCX-2.2", good).status == "Passed"
    result = _run("OOXML-DOCX-2.2", skip)
    assert result.status == "Failed"
    assert any("1 -> 3" in d or "1 → 3" in d for d in result.details)
    assert _run("OOXML-DOCX-2.2", first_deep).status == "Failed"


def test_rule_2_2_vacuous_pass_without_headings(tmp_path):
    none = make_docx(tmp_path / "n.docx", body_paragraphs=["No headings at all in here."])
    assert _run("OOXML-DOCX-2.2", none).status == "Passed"


def test_rule_2_3_no_orphan_intro_text(tmp_path):
    good = make_docx(tmp_path / "g.docx", headings=[("T", 0)],
                     body_paragraphs=["Body paragraph following the title."])
    bad = make_docx(tmp_path / "b.docx", headings=[("T", 0)],
                    body_paragraphs=["Intro before any heading.", "More body."],
                    body_first=True)
    empty = make_docx(tmp_path / "e.docx")
    assert _run("OOXML-DOCX-2.3", good).status == "Passed"
    result = _run("OOXML-DOCX-2.3", bad)
    assert result.status == "Failed" and result.details
    assert any("Intro before any heading." in d for d in result.details)
    assert _run("OOXML-DOCX-2.3", empty).status == "Passed"  # vacuous: nothing to mislead


def test_rule_2_1_outline_level_counts_as_heading(tmp_path):
    path = make_docx(tmp_path / "o.docx", outline_paragraphs=[("Outline heading", 0)])
    assert _run("OOXML-DOCX-2.1", path).status == "Passed"


def test_rule_2_2_outline_level_maps_to_level_plus_one(tmp_path):
    # outlineLvl 0 -> level 1: no skip
    ok = make_docx(tmp_path / "ok.docx", outline_paragraphs=[("Top", 0), ("Sub", 1)])
    assert _run("OOXML-DOCX-2.2", ok).status == "Passed"
    # outlineLvl 0 then 2 -> levels 1 then 3: skip
    skip = make_docx(tmp_path / "skip.docx", outline_paragraphs=[("Top", 0), ("Deep", 2)])
    result = _run("OOXML-DOCX-2.2", skip)
    assert result.status == "Failed"
    assert any("1 -> 3" in d for d in result.details)


def test_rule_3_1_inline_and_anchored_alt(tmp_path):
    good = make_docx(tmp_path / "g.docx", inline_images=1, image_alt="A chart")
    bad_inline = make_docx(tmp_path / "bi.docx", inline_images=1, image_alt=None)
    bad_anchored = make_docx(tmp_path / "ba.docx", inline_images=1, image_alt=None, anchored_images=True)
    assert _run("OOXML-DOCX-3.1", good).status == "Passed"
    r_inline = _run("OOXML-DOCX-3.1", bad_inline)
    assert r_inline.status == "Failed" and r_inline.rule_id == "docx-alt-text"
    # the baseline gap: anchored images must fail too (legacy check passed them)
    r_anchored = _run("OOXML-DOCX-3.1", bad_anchored)
    assert r_anchored.status == "Failed"
    assert any("anchored" in d for d in r_anchored.details)


def test_rule_3_2_placeholder_alt(tmp_path):
    good = make_docx(tmp_path / "g.docx", inline_images=1, image_alt="Campus map with entrances")
    bad = make_docx(tmp_path / "b.docx", inline_images=1, image_alt="image1.png")
    missing = make_docx(tmp_path / "m.docx", inline_images=1, image_alt=None)
    assert _run("OOXML-DOCX-3.2", good).status == "Passed"
    result = _run("OOXML-DOCX-3.2", bad)
    assert result.status == "Failed"
    assert any("image1.png" in d for d in result.details)
    # missing alt is 3.1's job; 3.2 passes vacuously
    assert _run("OOXML-DOCX-3.2", missing).status == "Passed"


def test_rule_4_1_table_header_marked(tmp_path):
    good = make_docx(tmp_path / "g.docx", tables=1, mark_table_headers=True)
    bad = make_docx(tmp_path / "b.docx", tables=1, mark_table_headers=False)
    assert _run("OOXML-DOCX-4.1", good).status == "Passed"
    result = _run("OOXML-DOCX-4.1", bad)
    assert result.status == "Failed" and result.rule_id == "docx-table-headers"
    assert any("table 1" in d for d in result.details)


def test_rule_4_2_no_merged_header_cells(tmp_path):
    good = make_docx(tmp_path / "g.docx", tables=1)
    bad = make_docx(tmp_path / "b.docx", tables=1, merge_header_cells=True)
    assert _run("OOXML-DOCX-4.2", good).status == "Passed"
    result = _run("OOXML-DOCX-4.2", bad)
    assert result.status == "Failed" and result.fixable is False


def test_rule_5_1_manual_bullets(tmp_path):
    good = make_docx(tmp_path / "g.docx", real_list_items=["alpha", "beta"])
    bad = make_docx(tmp_path / "b.docx", manual_bullets=["• first item", "- second item"])
    assert _run("OOXML-DOCX-5.1", good).status == "Passed"
    result = _run("OOXML-DOCX-5.1", bad)
    assert result.status == "Manual Check Needed"  # never a hard Fail (PRD §5/§10)
    assert len(result.details) == 2


def test_rule_6_1_link_text(tmp_path):
    good = make_docx(tmp_path / "g.docx", hyperlinks=[("District accessibility policy", "https://example.com/policy")])
    bare = make_docx(tmp_path / "b1.docx", hyperlinks=[("https://example.com", "https://example.com")])
    generic = make_docx(tmp_path / "b2.docx", hyperlinks=[("click here", "https://example.com")])
    assert _run("OOXML-DOCX-6.1", good).status == "Passed"
    assert _run("OOXML-DOCX-6.1", bare).status == "Failed"
    result = _run("OOXML-DOCX-6.1", generic)
    assert result.status == "Failed"
    assert any("click here" in d for d in result.details)


def test_rule_7_1_color_only_meaning(tmp_path):
    good = make_docx(tmp_path / "g.docx", color_paragraph="Deadlines are firm.")
    bad = make_docx(tmp_path / "b.docx", color_paragraph="Required fields are shown in red")
    plain = make_docx(tmp_path / "p.docx", body_paragraphs=["Items shown in red are required."])
    assert _run("OOXML-DOCX-7.1", good).status == "Passed"      # color without referential phrase
    result = _run("OOXML-DOCX-7.1", bad)
    assert result.status == "Manual Check Needed"               # color + phrase → flag, not fail
    assert _run("OOXML-DOCX-7.1", plain).status == "Passed"     # phrase without colored run


# --- Branch-coverage gap closers (AC1) --------------------------------------

import re
import shutil
import zipfile

import pytest


def _strip_first_docpr(path: Path) -> None:
    """Raw XML surgery: delete the first self-closing ``<wp:docPr .../>`` from
    word/document.xml, leaving a drawing container (wp:inline/wp:anchor) with
    no docPr child. Malformed per the OOXML schema but still well-formed XML
    and still readable by python-docx/ElementTree — exactly what's needed to
    exercise the "container found but no docPr" branch.
    """
    tmp = path.with_suffix(".tmp.docx")
    with zipfile.ZipFile(path) as src, zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as dst:
        for item in src.infolist():
            data = src.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                text = re.sub(r"<wp:docPr\b[^>]*/>", "", text, count=1)
                data = text.encode("utf-8")
            dst.writestr(item, data)
    shutil.move(str(tmp), str(path))


def test_rule_4_1_and_4_2_vacuous_pass_on_rowless_table(tmp_path):
    """A <w:tbl> with zero <w:tr> rows (python-docx add_table(rows=0, ...))
    exercises the `first_tr is None: continue` branch in both table rules —
    they must pass vacuously rather than raise or misreport.
    """
    from docx import Document

    path = tmp_path / "rowless.docx"
    doc = Document()
    table = doc.add_table(rows=0, cols=2)
    doc.save(str(path))

    # Sanity: confirm the fixture really has no w:tr (guards against a future
    # python-docx version changing add_table's zero-row behavior).
    with zipfile.ZipFile(path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
    tbl_start = xml.find("<w:tbl>")
    tbl_end = xml.find("</w:tbl>")
    assert "<w:tr" not in xml[tbl_start:tbl_end]

    assert _run("OOXML-DOCX-4.1", path).status == "Passed"
    assert _run("OOXML-DOCX-4.2", path).status == "Passed"


def test_rule_5_1_bullet_looking_text_with_real_numpr_not_flagged(tmp_path):
    path = make_docx(tmp_path / "real_dash_list.docx", real_list_items=["- dash item"])
    result = _run("OOXML-DOCX-5.1", path)
    assert result.status == "Passed"


def test_rule_6_1_empty_hyperlink_display_is_skipped(tmp_path):
    path = make_docx(tmp_path / "empty_link.docx", hyperlinks=[("", "https://example.com")])
    result = _run("OOXML-DOCX-6.1", path)
    assert result.status == "Passed"


def test_rule_3_1_drawing_without_docpr_is_vacuous(tmp_path):
    path = make_docx(tmp_path / "nodocpr.docx", inline_images=1, image_alt="alt text")
    _strip_first_docpr(path)
    result = _run("OOXML-DOCX-3.1", path)
    assert result.status == "Passed"


def test_run_all_delegates_non_docx(tmp_path):
    from project_remedy.models import FileType
    from project_remedy.office_checker import OfficeAccessibilityChecker
    from tests.unit.office_fixtures import make_pptx, make_xlsx

    pptx = make_pptx(tmp_path / "d.pptx", title="T")
    report = OfficeAccessibilityChecker(pptx).run_all()
    assert report.file_type == FileType.PPTX and report.results

    xlsx = make_xlsx(tmp_path / "b.xlsx", title="T")
    report = OfficeAccessibilityChecker(xlsx).run_all()
    assert report.file_type == FileType.XLSX and report.results

    with pytest.raises(ValueError, match="Unsupported Office checker type"):
        OfficeAccessibilityChecker(tmp_path / "x.doc", FileType.DOC).run_all()
